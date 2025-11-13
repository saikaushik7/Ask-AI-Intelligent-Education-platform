from flask import Blueprint, render_template, request, redirect, url_for, flash
from flask_login import login_required, current_user
import os, json, numpy as np
from werkzeug.utils import secure_filename

# OCR / Extraction
from pypdf import PdfReader
import docx
from pdf2image import convert_from_path
import pytesseract

# Tesseract path
pytesseract.pytesseract.tesseract_cmd = r"C:\\Program Files\\Tesseract-OCR\\tesseract.exe"

# DB models
from models import Document, db

# Summarizer (BART)
from transformers import BartTokenizer, BartForConditionalGeneration
import torch

# Gemini embeddings
import google.generativeai as genai
genai.configure(api_key=os.getenv("GEMINI_API_KEY"))
EMB_MODEL = "models/text-embedding-004"

# FAISS vector index
import faiss

docs_bp = Blueprint('docs_bp', __name__)

# ------------------- CONSTANTS -------------------
UPLOAD_FOLDER = 'uploads'
ALLOWED_EXTENSIONS = {'pdf', 'txt', 'docx'}
POPPLER_BIN = r"C:\poppler-25.07.0\Library\bin"
INDEX_DIR = os.path.join("instance", "indexes")
os.makedirs(INDEX_DIR, exist_ok=True)

# ------------------- BART -------------------
_DEVICE = "cpu"
_BART_TOKENIZER = BartTokenizer.from_pretrained("facebook/bart-large-cnn")
_BART_MODEL = BartForConditionalGeneration.from_pretrained("facebook/bart-large-cnn").to(_DEVICE)

def bart_summarize(text, target_words=50):
    approx_tokens = int(max(32, target_words * 1.35))
    inputs = _BART_TOKENIZER(text, return_tensors="pt", max_length=1024, truncation=True).to(_DEVICE)
    ids = _BART_MODEL.generate(
        inputs["input_ids"],
        num_beams=4,
        length_penalty=2.0,
        max_length=approx_tokens,
        min_length=max(20, approx_tokens // 2),
        no_repeat_ngram_size=3,
        early_stopping=True,
    )
    return _BART_TOKENIZER.decode(ids[0], skip_special_tokens=True)

# ------------------- Helpers -------------------
def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

def chunk_text_words(text: str, chunk_words=180, overlap_words=40):
    words = text.split()
    chunks = []
    i = 0
    stride = max(1, chunk_words - overlap_words)
    while i < len(words):
        chunk = words[i:i+chunk_words]
        if not chunk: break
        chunks.append(" ".join(chunk))
        i += stride
    return chunks

def embed_text_gemini(text: str) -> np.ndarray:
    resp = genai.embed_content(model=EMB_MODEL, content=text)
    v = np.array(resp["embedding"], dtype=np.float32)
    n = np.linalg.norm(v) + 1e-12
    return v / n

def build_faiss_index(doc_id: int, chunks: list[str]):
    if not chunks:
        return
    vecs = [embed_text_gemini(ch) for ch in chunks]
    X = np.vstack(vecs).astype(np.float32)
    d = X.shape[1]
    index = faiss.IndexFlatIP(d)
    index.add(X)

    base = os.path.abspath(INDEX_DIR)
    os.makedirs(base, exist_ok=True)
    faiss_path = os.path.join(base, f"{doc_id}.faiss")
    meta_path = os.path.join(base, f"{doc_id}.meta.json")

    faiss.write_index(index, faiss_path)
    with open(meta_path, "w", encoding="utf-8") as f:
        json.dump({"chunks": chunks}, f, ensure_ascii=False)

def extract_text_from_file(save_path: str, filename: str) -> str:
    text = ""
    if filename.lower().endswith(".pdf"):
        try:
            reader = PdfReader(save_path)
            for page in reader.pages:
                t = page.extract_text()
                if t:
                    text += t
        except Exception:
            pass
        if not text.strip():
            pages = convert_from_path(save_path, poppler_path=POPPLER_BIN)
            for img in pages:
                text += pytesseract.image_to_string(img)
    elif filename.lower().endswith(".docx"):
        d = docx.Document(save_path)
        for para in d.paragraphs:
            text += para.text + "\n"
    else:
        with open(save_path, "rb") as fh:
            raw = fh.read()
        text = raw.decode("utf-8", errors="ignore")
    return text

def get_latest_doc(user_id: int):
    return Document.query.filter_by(user_id=user_id).order_by(Document.id.desc()).first()

# ------------------- ROUTES -------------------
@docs_bp.route('/upload', methods=['GET', 'POST'])
@login_required
def upload():
    if request.method == 'POST':
        file = request.files.get('document')
        if not (file and allowed_file(file.filename)):
            flash("Please choose a valid file (.pdf / .docx / .txt).")
            return redirect(url_for('docs_bp.upload'))

        filename = secure_filename(file.filename)
        os.makedirs(UPLOAD_FOLDER, exist_ok=True)
        save_path = os.path.join(UPLOAD_FOLDER, filename)
        file.save(save_path)

        text = extract_text_from_file(save_path, filename)

        doc = Document(user_id=current_user.id, filename=filename, extracted_text=text)
        db.session.add(doc)
        db.session.commit()

        chunks = chunk_text_words(text, chunk_words=180, overlap_words=40)
        build_faiss_index(doc.id, chunks)

        flash("Document uploaded, extracted & indexed (FAISS).")
        return redirect(url_for('docs_bp.upload'))

    latest = get_latest_doc(current_user.id)
    extracted_text = latest.extracted_text if latest else None
    summary = latest.summary if (latest and latest.summary) else None
    total_words = len(extracted_text.split()) if extracted_text else 0
    summary_words = len(summary.split()) if summary else 0
    ratio = round((summary_words / total_words) * 100, 2) if total_words else 0.0

    latest_exists = bool(latest and latest.extracted_text and latest.extracted_text.strip())

    return render_template(
        "upload.html",
        latest_exists=latest_exists,
        extracted_text=extracted_text,
        summary=summary,
        total_words=total_words,
        summary_words=summary_words,
        ratio=ratio
    )

@docs_bp.route('/summarize', methods=['GET'])
@login_required
def summarize_latest():
    latest = get_latest_doc(current_user.id)
    if not latest:
        flash("Please upload a document first.")
        return redirect(url_for('docs_bp.upload'))

    text = latest.extracted_text or ""
    total_words = len(text.split())
    if total_words == 0:
        flash("No text extracted from the latest document.")
        return redirect(url_for('docs_bp.upload'))

    target_words = max(35, int(total_words * 0.35))
    raw = bart_summarize(text, target_words)
    raw_words = raw.split()
    final_summary = " ".join(raw_words[:target_words]) if len(raw_words) > target_words else raw

    latest.summary = final_summary
    db.session.commit()

    summary_words = len(final_summary.split())
    ratio = round((summary_words / total_words) * 100, 2)

    return render_template(
        "upload.html",
        latest_exists=True,
        extracted_text=text,
        summary=final_summary,
        total_words=total_words,
        summary_words=summary_words,
        ratio=ratio
    )

# ------------------- NEW ROUTES -------------------

@docs_bp.route('/summarize/<int:doc_id>')
@login_required
def summarize_specific(doc_id):
    """Summarize a specific document (for History page links)."""
    doc = Document.query.filter_by(id=doc_id, user_id=current_user.id).first()
    if not doc:
        flash("Document not found.", "danger")
        return redirect(url_for('docs_bp.history'))

    text = doc.extracted_text or ""
    if not text.strip():
        flash("This document has no extractable text.", "warning")
        return redirect(url_for('docs_bp.history'))

    total_words = len(text.split())
    target_words = max(35, int(total_words * 0.35))
    raw = bart_summarize(text, target_words)
    raw_words = raw.split()
    final_summary = " ".join(raw_words[:target_words]) if len(raw_words) > target_words else raw

    doc.summary = final_summary
    db.session.commit()

    summary_words = len(final_summary.split())
    ratio = round((summary_words / total_words) * 100, 2)

    return render_template(
        "upload.html",
        latest_exists=True,
        extracted_text=text,
        summary=final_summary,
        total_words=total_words,
        summary_words=summary_words,
        ratio=ratio
    )

@docs_bp.route('/history')
@login_required
def history():
    docs = Document.query.filter_by(user_id=current_user.id).order_by(Document.upload_date.desc()).all()
    return render_template("history.html", docs=docs)

@docs_bp.route('/delete_doc/<int:doc_id>')
@login_required
def delete_doc(doc_id):
    doc = Document.query.filter_by(id=doc_id, user_id=current_user.id).first()
    if not doc:
        flash("Document not found.", "danger")
        return redirect(url_for('docs_bp.history'))

    file_paths = [
        f"uploads/{doc.filename}",
        f"instance/indexes/{doc.id}.faiss",
        f"instance/indexes/{doc.id}.meta.json"
    ]
    for path in file_paths:
        abs_path = os.path.abspath(path)
        if os.path.exists(abs_path):
            os.remove(abs_path)

    db.session.delete(doc)
    db.session.commit()
    flash("Document deleted successfully.", "success")
    return redirect(url_for('docs_bp.history'))
